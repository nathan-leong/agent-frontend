"""
Usage:
    streamlit run streamlit_app.py
"""

import json
import os
import time
import logging
import uuid
import io
import re
import tempfile
import httpx
import boto3
import streamlit as st
from urllib.parse import quote
from docx import Document
from markdown_pdf import MarkdownPdf, Section
import fitz  # PyMuPDF

# logging.basicConfig(level=logging.DEBUG)  # root handler; optional if you only set botocore
# boto3.set_stream_logger("botocore", logging.DEBUG)
# optional: see lower-level HTTP details
# logging.getLogger("urllib3").setLevel(logging.DEBUG)
# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="NEXTGEN PDM Agent",
    page_icon="nextgen logo.png",
    layout="wide",
)

SAMPLE_DIR = os.path.dirname(os.path.abspath(__file__))
# ---------------------------------------------------------------------------
# Helper functions (self-contained, mirrors invoke.py patterns)
# ---------------------------------------------------------------------------


def _load_cognito_config() -> dict | None:
    """Load cognito_config.json from the sample directory."""
    path = os.path.join(SAMPLE_DIR, "cognito_config.json")
    if not os.path.exists(path):
        return None
    with open(path) as f:
        return json.load(f)


def _find_project_dir() -> str:
    """Find the agentcore project directory (subdirectory containing agentcore/)."""
    for entry in os.listdir(SAMPLE_DIR):
        candidate = os.path.join(SAMPLE_DIR, entry)
        if os.path.isdir(candidate) and os.path.isdir(os.path.join(candidate, "agentcore")):
            return candidate
    raise FileNotFoundError(
        "No agentcore project directory found. Run 'agentcore create' first."
    )


def _find_in_json(obj, key):
    """Recursively search for a key in nested JSON."""
    if isinstance(obj, dict):
        if key in obj:
            return obj[key]
        for v in obj.values():
            result = _find_in_json(v, key)
            if result:
                return result
    elif isinstance(obj, list):
        for item in obj:
            result = _find_in_json(item, key)
            if result:
                return result
    return None


def _resolve_agent_arn() -> str:
    """Read the deployed agent ARN from deployed-state.json.

    Searches for runtimeArn recursively to work across CLI versions.
    """
    project_dir = _find_project_dir()
    state_file = os.path.join(project_dir, "agentcore", ".cli", "deployed-state.json")
    if not os.path.exists(state_file):
        raise FileNotFoundError(
            "No deployed-state.json found. Run 'agentcore deploy -y' first."
        )
    with open(state_file) as f:
        state = json.load(f)
    arn = _find_in_json(state, "runtimeArn")
    if arn:
        return arn
    raise ValueError("No deployed agent found. Run 'agentcore deploy -y' first.")


def _get_bearer_token(config: dict) -> str:
    """Authenticate with Cognito and return an access token."""
    cognito = boto3.client("cognito-idp", region_name=config["region"])
    auth = cognito.initiate_auth(
        ClientId=config["client_id"],
        AuthFlow="USER_PASSWORD_AUTH",
        AuthParameters={
            "USERNAME": config["username"],
            "PASSWORD": config["password"],
        },
    )
    return auth["AuthenticationResult"]["AccessToken"]


def _get_aws_credentials_from_cognito(id_token: str, identity_pool_id: str, user_pool_id: str, region: str) -> dict:
    """Exchange Cognito ID token for AWS credentials via Identity Pool."""
    cognito_identity = boto3.client('cognito-identity', region_name=region)

    # Construct the login key for the user pool
    login_key = f'cognito-idp.{region}.amazonaws.com/{user_pool_id}'
    print(login_key)
    # Get Identity ID
    identity_response = cognito_identity.get_id(
        IdentityPoolId=identity_pool_id,
        Logins={
            login_key: id_token
        }
    )
    print('Identity Response')
    identity_id = identity_response['IdentityId']

    # Get credentials for the identity
    credentials_response = cognito_identity.get_credentials_for_identity(
        IdentityId=identity_id,
        Logins={
            login_key: id_token
        }
    )

    return credentials_response['Credentials']


def _list_s3_objects(bucket_name: str, prefix: str = "", delimiter: str = "") -> dict:
    """List objects in S3 bucket using temporary credentials."""
    if 'aws_credentials' not in st.session_state:
        return {'files': [], 'folders': []}

    creds = st.session_state.aws_credentials
    s3_client = boto3.client(
        's3',
        aws_access_key_id=creds['AccessKeyId'],
        aws_secret_access_key=creds['SecretKey'],
        aws_session_token=creds['SessionToken'],
        region_name='ap-southeast-2'
    )

    try:
        params = {'Bucket': bucket_name, 'Prefix': prefix}
        if delimiter:
            params['Delimiter'] = delimiter

        response = s3_client.list_objects_v2(**params)

        # Get folders (CommonPrefixes)
        folders = [p['Prefix'] for p in response.get('CommonPrefixes', [])]

        # Get files (Contents) that are not folders
        files = [obj for obj in response.get('Contents', []) if not obj['Key'].endswith('/')]

        return {'files': files, 'folders': folders}
    except Exception as e:
        st.error(f"Error listing objects: {e}")
        return {'files': [], 'folders': []}


def _upload_to_s3(bucket_name: str, file_obj, key: str) -> bool:
    """Upload file to S3 bucket."""
    if 'aws_credentials' not in st.session_state:
        return False

    creds = st.session_state.aws_credentials
    s3_client = boto3.client(
        's3',
        aws_access_key_id=creds['AccessKeyId'],
        aws_secret_access_key=creds['SecretKey'],
        aws_session_token=creds['SessionToken'],
        region_name='ap-southeast-2'
    )

    try:
        s3_client.upload_fileobj(file_obj, bucket_name, key)
        return True
    except Exception as e:
        st.error(f"Error uploading: {e}")
        return False


def _download_from_s3(bucket_name: str, key: str) -> bytes:
    """Download file from S3 bucket."""
    if 'aws_credentials' not in st.session_state:
        return None

    creds = st.session_state.aws_credentials
    s3_client = boto3.client(
        's3',
        aws_access_key_id=creds['AccessKeyId'],
        aws_secret_access_key=creds['SecretKey'],
        aws_session_token=creds['SessionToken'],
        region_name='ap-southeast-2'
    )

    try:
        response = s3_client.get_object(Bucket=bucket_name, Key=key)
        return response['Body'].read()
    except Exception as e:
        st.error(f"Error downloading: {e}")
        return None


def _delete_from_s3(bucket_name: str, key: str) -> bool:
    """Delete file from S3 bucket."""
    if 'aws_credentials' not in st.session_state:
        return False

    creds = st.session_state.aws_credentials
    s3_client = boto3.client(
        's3',
        aws_access_key_id=creds['AccessKeyId'],
        aws_secret_access_key=creds['SecretKey'],
        aws_session_token=creds['SessionToken'],
        region_name='ap-southeast-2'
    )

    try:
        s3_client.delete_object(Bucket=bucket_name, Key=key)
        return True
    except Exception as e:
        st.error(f"Error deleting: {e}")
        return False


def _start_kb_ingestion(kb_id: str, ds_id: str) -> str:
    """Start Bedrock Knowledge Base ingestion job."""
    if 'aws_credentials' not in st.session_state:
        return None

    creds = st.session_state.aws_credentials
    bedrock_agent = boto3.client(
        'bedrock-agent',
        aws_access_key_id=creds['AccessKeyId'],
        aws_secret_access_key=creds['SecretKey'],
        aws_session_token=creds['SessionToken'],
        region_name='ap-southeast-2'
    )

    try:
        response = bedrock_agent.start_ingestion_job(
            knowledgeBaseId=kb_id,
            dataSourceId=ds_id,
            description="Auto-sync after file operation"
        )
        job_id = response['ingestionJob']['ingestionJobId']
        return job_id
    except Exception as e:
        st.error(f"Error starting ingestion: {e}")
        return None


def _wait_for_ingestion_job(kb_id: str, ds_id: str, job_id: str, progress_bar, status_text) -> bool:
    """Monitor ingestion job status with progress updates."""
    if 'aws_credentials' not in st.session_state:
        return False

    creds = st.session_state.aws_credentials
    bedrock_agent = boto3.client(
        'bedrock-agent',
        aws_access_key_id=creds['AccessKeyId'],
        aws_secret_access_key=creds['SecretKey'],
        aws_session_token=creds['SessionToken'],
        region_name='ap-southeast-2'
    )

    try:
        max_checks = 60  # Max 10 minutes (60 * 10 seconds)
        check_count = 0

        while check_count < max_checks:
            response = bedrock_agent.get_ingestion_job(
                knowledgeBaseId=kb_id,
                dataSourceId=ds_id,
                ingestionJobId=job_id
            )

            status = response['ingestionJob']['status']
            check_count += 1

            # Update progress bar (0 to 100%)
            progress = min(check_count / max_checks, 0.95)  # Cap at 95% until complete
            progress_bar.progress(progress)
            status_text.text(f"Syncing Knowledge Base... Status: {status}")

            if status == 'COMPLETE':
                progress_bar.progress(1.0)
                status_text.text("✅ Knowledge Base sync completed!")
                return True
            elif status in ['FAILED', 'STOPPED']:
                status_text.text(f"⚠️ Sync ended with status: {status}")
                return False

            time.sleep(10)  # Wait 10 seconds between checks

        status_text.text("⚠️ Sync timeout - job is still running")
        return False

    except Exception as e:
        st.error(f"Error monitoring ingestion: {e}")
        return False


@st.dialog("📦 Knowledge Base Manager", width="large")
def knowledge_base_manager_modal():
    """Modal dialog for managing Knowledge Base S3 bucket - File Explorer style."""
    KB_BUCKET = "nextgen-aws-helper-agent-kb-692859930013-ap-southeast-2"
    KB_ROOT = "documents/"  # Root directory for the knowledge base
    KB_ID = "PGTFK15IS1"
    DS_ID = "SHZXLXGOK7"

    if not st.session_state.aws_credentials:
        st.warning("AWS credentials not available. Check cognito_config.json includes identity_pool_id and user_pool_id.")
        return

    # Initialize current path if not set or if it's outside KB_ROOT
    if 's3_current_prefix' not in st.session_state or not st.session_state.s3_current_prefix.startswith(KB_ROOT):
        st.session_state.s3_current_prefix = KB_ROOT

    current_path = st.session_state.s3_current_prefix

    # Breadcrumb navigation - clickable text
    relative_path = current_path[len(KB_ROOT):] if current_path.startswith(KB_ROOT) else current_path
    path_parts = [p for p in relative_path.rstrip('/').split('/') if p]

    # Display breadcrumb with buttons
    breadcrumb_cols = st.columns([0.1] + [0.15] * (len(path_parts) + 1) + [1])

    with breadcrumb_cols[0]:
        st.markdown("**Path:**")

    with breadcrumb_cols[1]:
        if st.button("🏠 documents", key="breadcrumb_root", use_container_width=True):
            st.session_state.s3_current_prefix = KB_ROOT
            st.rerun()

    # Add each path segment as a button
    for idx, part in enumerate(path_parts):
        with breadcrumb_cols[idx + 2]:
            path_value = KB_ROOT + '/'.join(path_parts[:idx + 1]) + '/'
            if st.button(f"/ {part}", key=f"breadcrumb_{idx}_{part}", use_container_width=True):
                st.session_state.s3_current_prefix = path_value
                st.rerun()

    # Action buttons
    col1, col2, col3 = st.columns([1, 1, 4])
    with col1:
        if st.button("⬆️ Upload File", key="show_upload", use_container_width=True):
            st.session_state.show_upload_form = True
    with col2:
        if st.button("📁 New Folder", key="show_new_folder", use_container_width=True):
            st.session_state.show_folder_form = True

    st.divider()

    # Show upload form if toggled
    if st.session_state.get('show_upload_form', False):
        with st.container():
            st.markdown("### Upload Files")
            upload_files = st.file_uploader("Choose files", key="kb_uploader_modal", accept_multiple_files=True)

            if upload_files:
                st.info(f"**Selected {len(upload_files)} file(s)**")

                # Show list of files to be uploaded
                for upload_file in upload_files:
                    st.text(f"📄 {upload_file.name} ({upload_file.size:,} bytes)")

                col1, col2 = st.columns([1, 1])
                with col1:
                    if st.button("✅ Upload All", key="do_upload", type="primary"):
                        progress_bar = st.progress(0)
                        success_count = 0

                        for idx, upload_file in enumerate(upload_files):
                            upload_key = f"{current_path}{upload_file.name}"
                            upload_file.seek(0)

                            if _upload_to_s3(KB_BUCKET, upload_file, upload_key):
                                success_count += 1

                            # Update progress
                            progress_bar.progress((idx + 1) / len(upload_files))

                        if success_count == len(upload_files):
                            st.success(f"✅ Uploaded all {success_count} file(s)")

                            # Start ingestion job
                            st.info("🔄 Starting Knowledge Base sync...")
                            job_id = _start_kb_ingestion(KB_ID, DS_ID)

                            if job_id:
                                sync_progress = st.progress(0)
                                sync_status = st.empty()

                                if _wait_for_ingestion_job(KB_ID, DS_ID, job_id, sync_progress, sync_status):
                                    st.success("✅ Files uploaded and Knowledge Base synced!")
                                else:
                                    st.warning("⚠️ Files uploaded but sync may still be in progress")
                        else:
                            st.warning(f"Uploaded {success_count} of {len(upload_files)} file(s)")

                        st.session_state.show_upload_form = False
                        time.sleep(2)
                        st.rerun()
                with col2:
                    if st.button("❌ Cancel", key="cancel_upload"):
                        st.session_state.show_upload_form = False
                        st.rerun()
            st.divider()

    # Show new folder form if toggled
    if st.session_state.get('show_folder_form', False):
        with st.container():
            st.markdown("### Create New Folder")
            folder_name = st.text_input("Folder name:", key="new_folder_name")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button("✅ Create", key="do_create_folder", type="primary"):
                    if folder_name:
                        # Create folder by uploading empty object with trailing /
                        folder_key = f"{current_path}{folder_name}/"
                        empty_file = io.BytesIO(b'')
                        if _upload_to_s3(KB_BUCKET, empty_file, folder_key):
                            st.success(f"Created folder: {folder_name}")

                            # Start ingestion job
                            st.info("🔄 Starting Knowledge Base sync...")
                            job_id = _start_kb_ingestion(KB_ID, DS_ID)

                            if job_id:
                                sync_progress = st.progress(0)
                                sync_status = st.empty()
                                _wait_for_ingestion_job(KB_ID, DS_ID, job_id, sync_progress, sync_status)

                            st.session_state.show_folder_form = False
                            time.sleep(1)
                            st.rerun()
            with col2:
                if st.button("❌ Cancel", key="cancel_folder"):
                    st.session_state.show_folder_form = False
                    st.rerun()
            st.divider()

    # List objects with delimiter to show folder structure
    result = _list_s3_objects(KB_BUCKET, current_path, delimiter='/')
    folders = result['folders']
    files = result['files']

    # Display folders and files
    with st.container(height=450):
        # Show folders first
        if folders:
            st.markdown("### 📁 Folders")
            for folder_prefix in folders:
                folder_name = folder_prefix.rstrip('/').split('/')[-1]

                col1, col2 = st.columns([5, 1])
                with col1:
                    if st.button(f"📁 {folder_name}", key=f"folder_{folder_prefix}", use_container_width=True):
                        st.session_state.s3_current_prefix = folder_prefix
                        st.rerun()
                with col2:
                    if st.button("🗑️", key=f"del_folder_{folder_prefix}", help="Delete folder"):
                        st.session_state.delete_confirm_key = folder_prefix
                        st.rerun()

                # Show confirmation dialog
                if st.session_state.delete_confirm_key == folder_prefix:
                    st.warning(f"⚠️ Delete folder **{folder_name}**? This will delete all contents.")
                    col_confirm1, col_confirm2 = st.columns([1, 1])
                    with col_confirm1:
                        if st.button("✅ Confirm Delete", key=f"confirm_del_folder_{folder_prefix}", type="primary"):
                            if _delete_from_s3(KB_BUCKET, folder_prefix):
                                st.success(f"Deleted folder: {folder_name}")

                                # Start ingestion job
                                st.info("🔄 Starting Knowledge Base sync...")
                                job_id = _start_kb_ingestion(KB_ID, DS_ID)

                                if job_id:
                                    sync_progress = st.progress(0)
                                    sync_status = st.empty()
                                    _wait_for_ingestion_job(KB_ID, DS_ID, job_id, sync_progress, sync_status)

                                st.session_state.delete_confirm_key = None
                                time.sleep(1)
                                st.rerun()
                    with col_confirm2:
                        if st.button("❌ Cancel", key=f"cancel_del_folder_{folder_prefix}"):
                            st.session_state.delete_confirm_key = None
                            st.rerun()

        # Show files
        if files:
            st.markdown("### 📄 Files")
            for obj in files:
                key = obj['Key']
                size = obj['Size']
                last_modified = obj['LastModified'].strftime('%Y-%m-%d %H:%M:%S')
                filename = key.split('/')[-1] if '/' in key else key

                col1, col2, col3 = st.columns([6, 0.5, 0.5])

                with col1:
                    st.markdown(f"**📄 {filename}**")
                    st.caption(f"{size:,} bytes • {last_modified}")

                with col2:
                    file_data = _download_from_s3(KB_BUCKET, key)
                    if file_data:
                        st.download_button(
                            label="⬇️",
                            data=file_data,
                            file_name=filename,
                            key=f"dl_{key}",
                            help="Download"
                        )

                with col3:
                    if st.button("🗑️", key=f"del_{key}", help="Delete"):
                        st.session_state.delete_confirm_key = key
                        st.rerun()

                # Show confirmation dialog
                if st.session_state.delete_confirm_key == key:
                    st.warning(f"⚠️ Delete file **{filename}**?")
                    col_confirm1, col_confirm2 = st.columns([1, 1])
                    with col_confirm1:
                        if st.button("✅ Confirm Delete", key=f"confirm_del_{key}", type="primary"):
                            if _delete_from_s3(KB_BUCKET, key):
                                st.success(f"Deleted {filename}")

                                # Start ingestion job
                                st.info("🔄 Starting Knowledge Base sync...")
                                job_id = _start_kb_ingestion(KB_ID, DS_ID)

                                if job_id:
                                    sync_progress = st.progress(0)
                                    sync_status = st.empty()
                                    _wait_for_ingestion_job(KB_ID, DS_ID, job_id, sync_progress, sync_status)

                                st.session_state.delete_confirm_key = None
                                time.sleep(1)
                                st.rerun()
                    with col_confirm2:
                        if st.button("❌ Cancel", key=f"cancel_del_{key}"):
                            st.session_state.delete_confirm_key = None
                            st.rerun()

                st.divider()

        if not folders and not files:
            st.info("📂 This folder is empty")


def _parse_event_stream(response: dict) -> str:
    """Extract text from the boto3 EventStream response."""
    parts: list[str] = []
    for event in response.get("response", []):
        raw = (
            event
            if isinstance(event, bytes)
            else event.get("chunk", {}).get("bytes", b"")
        )
        if raw:
            try:
                decoded = json.loads(raw.decode("utf-8"))
                if isinstance(decoded, str):
                    parts.append(decoded)
                elif isinstance(decoded, dict):
                    content = decoded.get("content", [])
                    for c in content:
                        if isinstance(c, dict) and c.get("type") == "text":
                            parts.append(c["text"])
                        elif isinstance(c, str):
                            parts.append(c)
                    if not content and "message" in decoded:
                        msg = decoded["message"]
                        if isinstance(msg, dict):
                            for c in msg.get("content", []):
                                if isinstance(c, dict) and c.get("type") == "text":
                                    parts.append(c["text"])
            except Exception:
                parts.append(raw.decode("utf-8"))
    raw = "\n".join(parts) if parts else "(no response)"
    return _sse_stream_to_plain_text(raw)


def _sse_stream_to_plain_text(s: str) -> str:
    """Turn SSE-style lines (`data: "chunk"\\n`) into plain concatenated text.

    Agent runtimes often return newline-delimited `data: "<json-string>"` chunks.
    If the payload does not look like SSE, return it unchanged.
    """
    lines = s.splitlines()
    if not any(line.strip().startswith("data:") for line in lines if line.strip()):
        return s
    out: list[str] = []
    for line in lines:
        line = line.strip()
        if not line.startswith("data:"):
            continue
        payload = line[5:].strip()
        if payload == "[DONE]":
            continue
        try:
            chunk = json.loads(payload)
            if isinstance(chunk, str):
                out.append(chunk)
            elif chunk is not None:
                out.append(str(chunk))
        except json.JSONDecodeError:
            out.append(payload)
    return "".join(out) if out else s

def _invoke_agent_streaming(
    agent_arn: str,
    region: str,
    prompt: str,
    bearer_token: str | None = None,
    session_id: str | None = None,
    on_chunk=None,  # callback for streaming UI
) -> dict:
    url = f"https://bedrock-agentcore.{region}.amazonaws.com/runtimes/{quote(agent_arn, safe='')}/invocations?qualifier=DEFAULT"

    headers = {
        "Content-Type": "application/json",
    }

    if bearer_token:
        headers["Authorization"] = f"Bearer {bearer_token}"

    if session_id:
        headers["X-Amzn-Bedrock-AgentCore-Runtime-Session-Id"] = session_id

    payload = {"prompt": prompt}

    t0 = time.time()
    full_text = []

    try:
        with httpx.Client(timeout=None) as client:
            with client.stream("POST", url, headers=headers, json=payload) as resp:
                resp.raise_for_status()

                # Stream line-by-line (SSE format)
                for line in resp.iter_lines():
                    if not line:
                        continue

                    line = line.strip()

                    if not line.startswith("data:"):
                        continue

                    data = line[5:].strip()

                    if data == "[DONE]":
                        break

                    try:
                        chunk = json.loads(data)
                        if isinstance(chunk, str):
                            text = chunk
                        else:
                            text = str(chunk)
                    except json.JSONDecodeError:
                        text = data

                    full_text.append(text)

                    # 🔥 STREAM OUT
                    if on_chunk:
                        on_chunk(text)

        elapsed = time.time() - t0

        return {
            "success": True,
            "text": "".join(full_text),
            "elapsed": elapsed,
            "error": None,
            "status_code": resp.status_code,
        }

    except Exception as exc:
        elapsed = time.time() - t0
        return {
            "success": False,
            "text": None,
            "elapsed": elapsed,
            "error": f"{type(exc).__name__}: {exc}",
            "status_code": getattr(exc, "response", None).status_code
            if hasattr(exc, "response") and exc.response
            else None,
        }


def _format_response(text: str) -> str:
    """Replace literal \\n sequences with real line breaks for display."""
    return text.replace("\\n", "\n")


def _truncate_arn(arn: str, max_len: int = 45) -> str:
    """Truncate an ARN for sidebar display."""
    if len(arn) <= max_len:
        return arn
    return arn[: max_len - 3] + "..."


def _create_markdown_file(question: str, answer: str, username: str) -> str:
    """Create a markdown string with the agent's response."""
    return answer


def _create_word_document(question: str, answer: str, username: str) -> io.BytesIO:
    """Create a Word document with the agent's response."""
    doc = Document()

    # Add answer - split by lines and add as paragraphs
    for line in answer.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _markdown_to_pdf(markdown_content: str) -> bytes:
    """Convert markdown content to PDF using markdown-pdf library with logo."""
    temp_pdf_path = tempfile.mktemp(suffix='.pdf')
    temp_pdf_with_logo = tempfile.mktemp(suffix='.pdf')

    try:
        # Add HTML spacing at the top for logo

        # Convert markdown to PDF - pass content directly
        pdf = MarkdownPdf()
        pdf.add_section(Section(markdown_content, toc=False))
        pdf.meta["title"] = "Document"
        pdf.save(temp_pdf_path)

        # Add logo and increase top margin using PyMuPDF
        logo_path = os.path.join(SAMPLE_DIR, "nextgen logo.png")
        doc = fitz.open(temp_pdf_path)

        # Create new document with shifted content
        new_doc = fitz.open()
        shift_down = 30  # 80 points = ~1.1 inches extra top margin

        for page_num, page in enumerate(doc):
            # Create new page with same dimensions
            new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)

            # Copy page content with vertical offset (shift down)
            new_page.show_pdf_page(
                new_page.rect + (0, shift_down, 0, 0),  # Shift down by 80 points
                doc,
                page_num
            )

            # Add logo to every page
            if os.path.exists(logo_path):
                logo_size = 40
                margin = 15
                page_width = new_page.rect.width
                x0 = page_width - logo_size - margin
                y0 = margin
                x1 = page_width - margin
                y1 = margin + logo_size

                rect = fitz.Rect(x0, y0, x1, y1)
                new_page.insert_image(rect, filename=logo_path, keep_proportion=True)

        # Save the modified document
        new_doc.save(temp_pdf_with_logo)
        new_doc.close()
        doc.close()

        # Read the modified PDF
        with open(temp_pdf_with_logo, 'rb') as f:
            pdf_bytes = f.read()

        return pdf_bytes
    finally:
        # Clean up temp files
        try:
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            if os.path.exists(temp_pdf_with_logo):
                os.unlink(temp_pdf_with_logo)
        except:
            pass


# ---------------------------------------------------------------------------
# Session state initialisation
# ---------------------------------------------------------------------------
for key, default in {
    "logged_in": False,
    "jwt_token": None,
    "id_token": None,
    "username": "",
    "chat_history": [],
    "agent_arn": None,
    "arn_error": None,
    "region": "us-east-1",
    "last_request": None,
    "login_error": None,
    "session_id": None,
    "aws_credentials": None,
    "user_pool_id": None,
    "s3_current_prefix": "documents/",
    "show_upload_form": False,
    "show_folder_form": False,
    "kb_manager_open": False,
    "delete_confirm_key": None,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# ---------------------------------------------------------------------------
# Load config
# ---------------------------------------------------------------------------
config = _load_cognito_config()

# ---------------------------------------------------------------------------
# Custom CSS
# ---------------------------------------------------------------------------
st.markdown(
    """
    <style>
    /* Hide sidebar toggle when not logged in */
    .login-page [data-testid="collapsedControl"] { display: none; }

    /* Login card styling */
    .login-card {
        padding: 2rem 0;
    }
    .login-header {
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .login-subtitle {
        text-align: center;
        color: #6b7280;
        font-size: 0.95rem;
        margin-bottom: 1.5rem;
    }
    .login-desc {
        text-align: center;
        color: #9ca3af;
        font-size: 0.85rem;
        margin-bottom: 2rem;
        line-height: 1.5;
    }

    /* Sidebar signed-in badge */
    .signed-in-badge {
        background: #065f46;
        color: #d1fae5;
        padding: 0.4rem 0.75rem;
        border-radius: 0.5rem;
        font-size: 0.85rem;
        font-weight: 500;
        text-align: center;
        margin-bottom: 0.5rem;
    }
    .sidebar-meta {
        color: #9ca3af;
        font-size: 0.75rem;
        word-break: break-all;
        margin-bottom: 0.25rem;
    }

    /* Preset buttons row */
    .stButton > button {
        font-size: 0.85rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# =========================================================================
# SCREEN 1: LOGIN (full page, centered)
# =========================================================================
if not st.session_state.logged_in:
    # Hide sidebar on login page
    st.markdown(
        "<style>[data-testid='stSidebar'] { display: none; }</style>",
        unsafe_allow_html=True,
    )

    # Vertical spacer
    st.markdown("")
    st.markdown("")

    # Center column
    _, center, _ = st.columns([1, 2, 1])

    with center:
        # Display logo
        logo_col1, logo_col2, logo_col3 = st.columns([1, 1, 1])
        with logo_col2:
            st.image("nextgen logo.png", width=150)

        st.markdown(
            "<h1 class='login-header'>NEXTGEN Partner Development Manager Agent</h1>",
            unsafe_allow_html=True,
        )

        if not config:
            st.error(
                "**cognito_config.json not found.** "
                "Run `python setup_cognito.py` before using this app."
            )
            st.stop()

        # Show any previous login error
        if st.session_state.login_error:
            st.error(st.session_state.login_error)

        with st.form("login_form"):
            username = st.text_input(
                "Username",
                value=config.get("username", ""),
            )
            password = st.text_input(
                "Password",
                value=config.get("password", ""),
                type="password",
            )
            sign_in = st.form_submit_button("Sign In", use_container_width=True)

        if sign_in:
            if not username or not password:
                st.session_state.login_error = "Username and password are required."
                st.rerun()
            else:
                login_config = {**config, "username": username, "password": password}
                try:
                    with st.spinner("Authenticating with Cognito..."):
                        cognito = boto3.client("cognito-idp", region_name=config["region"])
                        auth = cognito.initiate_auth(
                            ClientId=config["client_id"],
                            AuthFlow="USER_PASSWORD_AUTH",
                            AuthParameters={
                                "USERNAME": username,
                                "PASSWORD": password,
                            },
                        )
                        access_token = auth["AuthenticationResult"]["AccessToken"]
                        id_token = auth["AuthenticationResult"]["IdToken"]

                    # Exchange ID token for AWS credentials if identity_pool_id exists
                    aws_credentials = None
                    if config.get("identity_pool_id") and config.get("user_pool_id"):
                        with st.spinner("Getting AWS credentials..."):
                            try:
                                aws_credentials = _get_aws_credentials_from_cognito(
                                    id_token,
                                    config["identity_pool_id"],
                                    config["user_pool_id"],
                                    config["region"]
                                )
                                st.session_state.user_pool_id = config["user_pool_id"]
                            except Exception as exc:
                                st.warning(f"Could not get AWS credentials: {exc}")

                    # Resolve agent ARN right away
                    with st.spinner("Resolving agent ARN..."):
                        try:
                            agent_arn = _resolve_agent_arn()
                        except Exception as exc:
                            agent_arn = None
                            st.session_state.arn_error = str(exc)

                    # Commit to session state
                    st.session_state.jwt_token = access_token
                    st.session_state.id_token = id_token
                    st.session_state.bearer_input = access_token  # pre-fill the token field
                    st.session_state.username = username
                    st.session_state.logged_in = True
                    st.session_state.agent_arn = agent_arn
                    st.session_state.region = config.get("region", "us-east-1")
                    st.session_state.session_id = str(uuid.uuid4())  # Generate session ID
                    st.session_state.aws_credentials = aws_credentials
                    st.session_state.login_error = None
                    st.rerun()

                except Exception as exc:
                    st.session_state.login_error = f"Login failed: {exc}"
                    st.rerun()

    st.stop()


# =========================================================================
# SCREEN 2: DASHBOARD (after login)
# =========================================================================

# --- Sidebar ---
with st.sidebar:
    st.markdown(
        f"<div class='signed-in-badge'>Signed in as {st.session_state.username}</div>",
        unsafe_allow_html=True,
    )

    if st.session_state.agent_arn:
        st.markdown(
            f"<p class='sidebar-meta'><b>Agent:</b> {_truncate_arn(st.session_state.agent_arn)}</p>",
            unsafe_allow_html=True,
        )
    elif st.session_state.arn_error:
        st.error(st.session_state.arn_error, icon="\u26a0\ufe0f")
    else:
        st.warning("Agent ARN not resolved.")

    st.markdown(
        f"<p class='sidebar-meta'><b>Region:</b> {st.session_state.region}</p>",
        unsafe_allow_html=True,
    )

    if st.session_state.session_id:
        st.markdown(
            f"<p class='sidebar-meta'><b>Session ID:</b> {st.session_state.session_id[:8]}...</p>",
            unsafe_allow_html=True,
        )

    if st.button("Sign Out", use_container_width=True):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

    st.divider()

    # --- Tools Section ---
    st.markdown("### 🛠️ Tools")

    # Tool 1: Copy Bearer Token
    with st.expander("📋 Copy Bearer Token", expanded=False):
        if st.session_state.jwt_token:
            st.code(st.session_state.jwt_token, language=None)
            st.caption("Copy the token above to use in other applications.")
        else:
            st.info("No token available")

    # Tool 2: Markdown to PDF Converter
    with st.expander("📄 Markdown to PDF Converter", expanded=False):
        uploaded_file = st.file_uploader(
            "Upload markdown file",
            type=['md', 'markdown'],
            key="md_uploader",
            label_visibility="collapsed"
        )

        if uploaded_file is not None:
            # Read the markdown content
            markdown_content = uploaded_file.read().decode('utf-8')

            # Convert to PDF
            pdf_bytes = _markdown_to_pdf(markdown_content)

            # Download button
            st.download_button(
                label="⬇️ Download PDF",
                data=pdf_bytes,
                file_name=f"{uploaded_file.name.rsplit('.', 1)[0]}.pdf",
                mime="application/pdf",
                key="md_to_pdf_download",
                use_container_width=True
            )

    # Tool 3: Knowledge Base S3 Manager
    if st.button("📦 Knowledge Base Manager", use_container_width=True, key="open_kb_manager"):
        st.session_state.kb_manager_open = True
        st.session_state.s3_current_prefix = "documents/"  # Reset to root on open

# Open KB Manager modal if flag is set
if st.session_state.kb_manager_open:
    knowledge_base_manager_modal()

# --- Main area ---
# Display logo in the app
logo_left, logo_right = st.columns([1, 9])
with logo_left:
    st.image("nextgen logo.png", width=50)
with logo_right:
    st.markdown("#### NEXTGEN Partner Development Manager Agent")

# st.caption("Clear the Bearer Token in the sidebar to see a 403 rejection. The agent retrieves the API key from AgentCore Identity — never hardcoded. The agent is deployed in AWS using AgentCore.")

# Chat history
for i, msg in enumerate(st.session_state.chat_history):
    with st.chat_message(msg["role"]):
        st.markdown(_format_response(msg["content"]))

    # Add download buttons for assistant responses (outside chat_message)
    if msg["role"] == "assistant" and msg["content"]:
        # Find the corresponding user question
        user_question = ""
        if i > 0 and st.session_state.chat_history[i-1]["role"] == "user":
            user_question = st.session_state.chat_history[i-1]["content"]

        # Create markdown content
        markdown_content = _create_markdown_file(
            user_question,
            msg["content"],
            st.session_state.username
        )

        # Create Word document
        doc_buffer = _create_word_document(
            user_question,
            msg["content"],
            st.session_state.username
        )

        # Add download buttons with small styling
        col1, col2, col3, col4 = st.columns([1, 2, 2, 5])
        with col2:
            st.download_button(
                label="📄 Markdown",
                data=markdown_content,
                file_name=f"agent_response_{time.strftime('%Y%m%d_%H%M%S')}.md",
                mime="text/markdown",
                key=f"download_md_{i}",
                use_container_width=True
            )
        with col3:
            st.download_button(
                label="📘 Word",
                data=doc_buffer,
                file_name=f"agent_response_{time.strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"download_docx_{i}",
                use_container_width=True
            )

# Preset buttons
prompt_to_send = None

if len(st.session_state.chat_history) == 0:
    presets = [
        "What can you do?",
        "Show me all opportunities created this month.",
        "How much discounts do software partners receive through the distribution program discounts?",
        "What are the requirements for a partner to become Advanced Tier?"
    ]

    preset_cols = st.columns(len(presets))

    for i, preset in enumerate(presets):
        with preset_cols[i]:
            if st.button(preset, key=f"preset_{i}", use_container_width=True):
                prompt_to_send = preset

# Chat input
user_input = st.chat_input("Ask the agent...")
if user_input:
    prompt_to_send = user_input

# Send prompt
if prompt_to_send:
    if not st.session_state.agent_arn:
        st.error("Agent ARN not resolved. Deploy the agent first.")
    else:
        st.session_state.chat_history.append({
            "role": "user",
            "content": prompt_to_send,
        })
        with st.chat_message("user"):
            st.markdown(prompt_to_send)

        with st.chat_message("assistant"):
            placeholder = st.empty()

            streamed_text = [""]
            first_chunk_received = [False]

            spinner_container = st.empty()
            with spinner_container:
                st.spinner("Invoking agent...")

            # 🔥 Initialize message in chat history BEFORE streaming
            st.session_state.chat_history.append({
                "role": "assistant",
                "content": "",
            })
            message_index = len(st.session_state.chat_history) - 1

            def handle_chunk(chunk):
                if not first_chunk_received[0]:
                    first_chunk_received[0] = True
                    spinner_container.empty()

                streamed_text[0] += chunk

                # 🔥 Persist LIVE into session state
                st.session_state.chat_history[message_index]["content"] = streamed_text[0]

                placeholder.markdown(_format_response(streamed_text[0]))

            result = _invoke_agent_streaming(
                st.session_state.agent_arn,
                st.session_state.region,
                prompt_to_send,
                bearer_token=st.session_state.get("bearer_input", "").strip() or None,
                session_id=st.session_state.get("session_id"),
                on_chunk=handle_chunk,
            )

            if not first_chunk_received[0]:
                spinner_container.empty()

            # Rerun to show download buttons
            st.rerun()
# Last request details (collapsed)
if st.session_state.last_request:
    with st.expander("Last request details", expanded=False):
        req = st.session_state.last_request
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Status", "Success" if req.get("success") else "Failed")
        with col2:
            st.metric("Response Time", f"{req.get('elapsed', 0):.2f}s")
        with col3:
            st.metric("HTTP Status", str(req.get("status_code", "N/A")))
        st.code(f"Authorization: {req.get('auth', 'N/A')}", language=None)
        if req.get("error"):
            st.error(req["error"])
        if req.get("text"):
            st.code(req["text"], language=None)